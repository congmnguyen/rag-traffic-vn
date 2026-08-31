from __future__ import annotations

import hashlib
import json
import os
import re
import shutil
import subprocess
import tempfile
import unicodedata
from collections.abc import Iterable
from dataclasses import dataclass, asdict
from pathlib import Path


SUPPORTED_EXTENSIONS = {".doc", ".docx", ".pdf"}
ARTICLE_RE = re.compile(r"(?im)^\s*(Điều\s+\d+[a-zđ]?(?:\s*[.:–-]|\s+).*)$")
EXPLICIT_DOCUMENT_ID_RE = re.compile(
    r"(?im)^\s*Số\s*:\s*(\d{1,3})\s*/\s*(\d{4})\s*/\s*"
    r"(TTLT|TT|NĐ|ND|QĐ|QD)\s*[-–]\s*([A-ZĐ]+(?:[-–][A-ZĐ]+)*)"
)
FILENAME_DOCUMENT_ID_RE = re.compile(
    r"(?i)(\d{1,3})[. _/-]+(\d{4})[. _/-]+(TTLT|TT|NĐ|ND|QĐ|QD)"
    r"[. _/-]+([A-ZĐ]+(?:[. _-]+[A-ZĐ]+)*)"
)
DATE_PATTERNS = (
    re.compile(r"(?i)(ngày\s+\d{1,2}\s+tháng\s+\d{1,2}\s+năm\s+\d{4})"),
    re.compile(r"\b(\d{1,2}/\d{1,2}/\d{4})\b"),
)


@dataclass(frozen=True)
class Chunk:
    chunk_id: str
    document_id: str
    document_id_source: str
    title: str
    issuing_agency: str
    issued_date: str
    article: str
    source_file: str
    source_sha256: str
    content_sha256: str
    content: str

    def to_dict(self) -> dict[str, str]:
        return asdict(self)


@dataclass
class IngestionReport:
    files_seen: int = 0
    files_processed: int = 0
    files_skipped: int = 0
    chunks_written: int = 0
    exact_duplicates_removed: int = 0
    errors: list[dict[str, str]] | None = None
    exclusions: list[dict[str, str]] | None = None

    def __post_init__(self) -> None:
        if self.errors is None:
            self.errors = []
        if self.exclusions is None:
            self.exclusions = []


def normalize_text(text: str) -> str:
    text = unicodedata.normalize("NFC", text.replace("\x00", " "))
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r" *\n *", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def normalized_content_hash(text: str) -> str:
    canonical = re.sub(r"\s+", " ", normalize_text(text)).casefold()
    return hashlib.sha256(canonical.encode("utf-8")).hexdigest()


def is_draft_source(path: Path) -> bool:
    ascii_name = unicodedata.normalize("NFKD", path.name).encode("ascii", "ignore").decode()
    return bool(re.search(r"(?i)(?:^|[ _.-])du[ _.-]*thao(?:[ _.-]|$)", ascii_name))


def file_sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for block in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def extract_docx(path: Path) -> str:
    from docx import Document
    from docx.table import Table

    document = Document(path)
    blocks: list[str] = []
    # `document.paragraphs` followed by `document.tables` destroys the original
    # order and can detach a legal article heading from the table containing its
    # values. python-docx 1.2 exposes the blocks in document order.
    for block in document.iter_inner_content():
        if isinstance(block, Table):
            for row in block.rows:
                blocks.append(" | ".join(cell.text for cell in row.cells))
        else:
            blocks.append(block.text)
    return normalize_text("\n".join(blocks))


def extract_pdf(path: Path) -> str:
    from pdfminer.high_level import extract_text

    text = normalize_text(extract_text(str(path)))
    if len(text) >= 200:
        return text

    # OCR is intentionally a fallback because it is much slower than reading a
    # searchable PDF and can introduce errors in article numbers.
    from pdf2image import convert_from_path
    import pytesseract

    configure_tesseract(pytesseract)
    pages = convert_from_path(str(path), dpi=250)
    return normalize_text("\n".join(pytesseract.image_to_string(p, lang="vie") for p in pages))


def configure_tesseract(pytesseract_module: object) -> Path:
    configured = os.getenv("RAG_TESSERACT_CMD")
    system_binary = shutil.which("tesseract")
    local_root = Path(
        os.getenv("RAG_TESSERACT_ROOT", str(Path.home() / ".local/opt/tesseract"))
    ).expanduser()
    binary = (
        Path(configured).expanduser()
        if configured
        else Path(system_binary)
        if system_binary
        else local_root / "usr/bin/tesseract"
    )
    if not binary.is_file():
        raise RuntimeError(
            "Không tìm thấy Tesseract. Cài tesseract-ocr-vie hoặc đặt RAG_TESSERACT_CMD."
        )

    # A sudo-less Debian extraction also needs its private shared library and
    # tessdata paths passed to the child process started by pytesseract.
    local_lib = local_root / "usr/lib/x86_64-linux-gnu"
    local_data = local_root / "usr/share/tesseract-ocr/5/tessdata"
    if binary == local_root / "usr/bin/tesseract":
        current_library_path = os.environ.get("LD_LIBRARY_PATH", "")
        library_parts = [part for part in current_library_path.split(":") if part]
        if str(local_lib) not in library_parts:
            os.environ["LD_LIBRARY_PATH"] = ":".join([str(local_lib), *library_parts])
        os.environ.setdefault("TESSDATA_PREFIX", str(local_data))

    pytesseract_module.pytesseract.tesseract_cmd = str(binary)  # type: ignore[attr-defined]
    return binary


def extract_doc(path: Path) -> str:
    soffice = shutil.which("soffice")
    if not soffice:
        raise RuntimeError("Cần LibreOffice/soffice để đọc định dạng .doc cũ")
    with tempfile.TemporaryDirectory(prefix="rag-traffic-doc-") as temp_dir:
        result = subprocess.run(
            [soffice, "--headless", "--convert-to", "txt:Text", "--outdir", temp_dir, str(path)],
            check=False,
            capture_output=True,
            text=True,
            timeout=120,
        )
        output = Path(temp_dir) / f"{path.stem}.txt"
        if result.returncode != 0 or not output.exists():
            message = result.stderr.strip() or result.stdout.strip() or "không rõ lỗi"
            raise RuntimeError(f"Không chuyển đổi được .doc: {message}")
        return normalize_text(output.read_text(encoding="utf-8", errors="replace"))


def extract_text(path: Path) -> str:
    extension = path.suffix.lower()
    if extension == ".docx":
        return extract_docx(path)
    if extension == ".pdf":
        return extract_pdf(path)
    if extension == ".doc":
        return extract_doc(path)
    raise ValueError(f"Không hỗ trợ định dạng {extension}")


def _canonical_document_id(match: re.Match[str]) -> str:
    number, year, kind, agency = match.groups()
    kind = "NĐ" if kind.upper() == "ND" else "QĐ" if kind.upper() == "QD" else kind.upper()
    agency = re.sub(r"[. _–]+", "-", agency.upper()).strip("-")
    return f"{number}/{year}/{kind}-{agency}"


def extract_document_identity(text: str, path: Path) -> tuple[str, str]:
    header_match = EXPLICIT_DOCUMENT_ID_RE.search(text[:8000])
    if header_match:
        return _canonical_document_id(header_match), "header"

    filename = re.sub(r"(?i)^VanBanGoc[_ -]*", "", path.stem)
    filename_match = FILENAME_DOCUMENT_ID_RE.search(filename)
    if filename_match:
        return _canonical_document_id(filename_match), "filename"

    return (filename.strip() or path.stem), "filename_fallback"


def extract_title(text: str, fallback: str) -> str:
    lines = [line.strip() for line in text[:12000].splitlines() if line.strip()]
    for index, line in enumerate(lines):
        if re.search(r"(?i)\b(THÔNG TƯ|NGHỊ ĐỊNH|QUYẾT ĐỊNH|QUY CHUẨN)\b", line):
            candidates = lines[index : index + 6]
            return " — ".join(candidates[:2])[:500]
    return fallback


def extract_agency(text: str) -> str:
    for line in text[:5000].splitlines():
        line = line.strip()
        if re.match(r"(?i)^(BỘ|CHÍNH PHỦ|LIÊN BỘ|ỦY BAN)\b", line):
            return line[:250]
    return ""


def extract_date(text: str) -> str:
    for pattern in DATE_PATTERNS:
        match = pattern.search(text[:10000])
        if match:
            return match.group(1)
    return ""


def split_sections(text: str) -> list[tuple[str, str]]:
    matches = list(ARTICLE_RE.finditer(text))
    if not matches:
        return [("", text)]
    sections: list[tuple[str, str]] = []
    preamble = text[: matches[0].start()].strip()
    if preamble:
        sections.append(("Phần mở đầu", preamble))
    for index, match in enumerate(matches):
        end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
        sections.append((match.group(1).strip(), text[match.start() : end].strip()))
    return sections


def chunk_section(text: str, max_chars: int = 1800, overlap_chars: int = 180) -> list[str]:
    text = text.strip()
    if len(text) <= max_chars:
        return [text] if text else []
    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(start + max_chars, len(text))
        if end < len(text):
            boundary = max(text.rfind("\n", start + max_chars // 2, end), text.rfind(". ", start + max_chars // 2, end))
            if boundary > start:
                end = boundary + 1
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(text):
            break
        start = max(end - overlap_chars, start + 1)
    return chunks


def iter_source_files(source_dir: Path) -> Iterable[Path]:
    # DOCX first: it normally has cleaner text than a PDF copy of the same law.
    priority = {".docx": 0, ".doc": 1, ".pdf": 2}
    files = (
        path
        for path in source_dir.rglob("*")
        if path.is_file()
        and path.suffix.lower() in SUPPORTED_EXTENSIONS
        and not path.name.startswith((".~lock", "~$"))
    )
    return sorted(files, key=lambda p: (priority[p.suffix.lower()], str(p).casefold()))


def ingest(source_dir: Path, output_file: Path) -> IngestionReport:
    if not source_dir.is_dir():
        raise FileNotFoundError(f"Không tìm thấy thư mục dữ liệu: {source_dir}")

    report = IngestionReport()
    chunks: list[Chunk] = []
    seen_content: set[str] = set()
    files = list(iter_source_files(source_dir))
    report.files_seen = len(files)

    for path in files:
        if is_draft_source(path):
            report.files_skipped += 1
            assert report.exclusions is not None
            report.exclusions.append(
                {"source_file": str(path), "reason": "Loại dự thảo khỏi index mặc định"}
            )
            continue
        try:
            text = extract_text(path)
            if not text:
                raise ValueError("không trích xuất được văn bản")
            source_hash = file_sha256(path)
            document_id, document_id_source = extract_document_identity(text, path)
            title = extract_title(text, path.stem)
            agency = extract_agency(text)
            issued_date = extract_date(text)
            relative_source = str(path.relative_to(source_dir))

            for article, section in split_sections(text):
                for part_number, content in enumerate(chunk_section(section), start=1):
                    content_hash = normalized_content_hash(content)
                    if content_hash in seen_content:
                        report.exact_duplicates_removed += 1
                        continue
                    seen_content.add(content_hash)
                    identity = f"{relative_source}\0{article}\0{part_number}\0{content_hash}"
                    chunks.append(
                        Chunk(
                            chunk_id=hashlib.sha256(identity.encode("utf-8")).hexdigest()[:24],
                            document_id=document_id,
                            document_id_source=document_id_source,
                            title=title,
                            issuing_agency=agency,
                            issued_date=issued_date,
                            article=article,
                            source_file=relative_source,
                            source_sha256=source_hash,
                            content_sha256=content_hash,
                            content=content,
                        )
                    )
            report.files_processed += 1
        except Exception as exc:  # keep a full report instead of losing a long batch
            report.files_skipped += 1
            assert report.errors is not None
            report.errors.append({"source_file": str(path), "error": str(exc)})

    output_file.parent.mkdir(parents=True, exist_ok=True)
    temporary = output_file.with_suffix(output_file.suffix + ".tmp")
    with temporary.open("w", encoding="utf-8") as stream:
        for chunk in chunks:
            stream.write(json.dumps(chunk.to_dict(), ensure_ascii=False) + "\n")
    temporary.replace(output_file)
    report.chunks_written = len(chunks)
    return report
