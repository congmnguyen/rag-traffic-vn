from pathlib import Path
from types import SimpleNamespace

import rag_traffic.ingest as ingestion


def test_split_sections_keeps_preamble_and_articles() -> None:
    text = "THÔNG TƯ\nPhần dẫn nhập\nĐiều 1. Phạm vi\nNội dung một\nĐiều 2. Đối tượng\nNội dung hai"
    sections = ingestion.split_sections(text)
    assert [name for name, _ in sections] == [
        "Phần mở đầu",
        "Điều 1. Phạm vi",
        "Điều 2. Đối tượng",
    ]
    assert sections[1][1].startswith("Điều 1")


def test_chunk_section_is_bounded_and_overlaps() -> None:
    text = " ".join(f"từ-{number}." for number in range(500))
    chunks = ingestion.chunk_section(text, max_chars=300, overlap_chars=30)
    assert len(chunks) > 1
    assert all(len(chunk) <= 300 for chunk in chunks)


def test_draft_source_detection_handles_vietnamese_and_ascii() -> None:
    assert ingestion.is_draft_source(Path("Dự thảo thông tư.docx"))
    assert ingestion.is_draft_source(Path("Du thao 2 sua doi.docx"))
    assert not ingestion.is_draft_source(Path("Thông tư 02.docx"))


def test_document_id_prefers_filename_over_referenced_law() -> None:
    text = "Căn cứ Nghị định số 12/2017/NĐ-CP của Chính phủ"
    identity, source = ingestion.extract_document_identity(
        text, Path("TT.31.2019.TT.BGTVT.docx")
    )
    assert identity == "31/2019/TT-BGTVT"
    assert source == "filename"


def test_document_id_prefers_explicit_header() -> None:
    text = "BỘ GIAO THÔNG VẬN TẢI\nSố: 06/2024/TT-BGTVT\nCăn cứ 12/2017/NĐ-CP"
    identity, source = ingestion.extract_document_identity(text, Path("ban-scan.pdf"))
    assert identity == "06/2024/TT-BGTVT"
    assert source == "header"


def test_docx_extraction_preserves_paragraph_table_order(tmp_path: Path) -> None:
    from docx import Document

    path = tmp_path / "ordered.docx"
    document = Document()
    document.add_paragraph("Điều 1. Bảng tốc độ")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Xe mô tô"
    table.cell(0, 1).text = "60 km/h"
    document.add_paragraph("Điều 2. Quy định tiếp theo")
    document.save(path)

    text = ingestion.extract_docx(path)
    assert text.index("Điều 1") < text.index("Xe mô tô | 60 km/h") < text.index("Điều 2")


def test_configure_tesseract_detects_user_local_install(tmp_path: Path, monkeypatch) -> None:
    root = tmp_path / "tesseract"
    binary = root / "usr/bin/tesseract"
    binary.parent.mkdir(parents=True)
    binary.write_bytes(b"binary")
    monkeypatch.setenv("RAG_TESSERACT_ROOT", str(root))
    monkeypatch.delenv("RAG_TESSERACT_CMD", raising=False)
    monkeypatch.setattr(ingestion.shutil, "which", lambda _: None)
    fake_module = SimpleNamespace(pytesseract=SimpleNamespace(tesseract_cmd=None))

    detected = ingestion.configure_tesseract(fake_module)

    assert detected == binary
    assert fake_module.pytesseract.tesseract_cmd == str(binary)


def test_ingest_removes_exact_duplicates_and_has_unique_ids(tmp_path: Path, monkeypatch) -> None:
    source = tmp_path / "source"
    source.mkdir()
    (source / "a.docx").write_bytes(b"a")
    (source / "b.pdf").write_bytes(b"b")
    text = "THÔNG TƯ\nQuy định thử nghiệm\nĐiều 1. Phạm vi\nCùng một nội dung."
    monkeypatch.setattr(ingestion, "extract_text", lambda _: text)

    output = tmp_path / "chunks.jsonl"
    report = ingestion.ingest(source, output)
    chunks = ingestion.json.loads("[" + ",".join(output.read_text().splitlines()) + "]")

    assert report.files_processed == 2
    assert report.exact_duplicates_removed == 2
    assert len(chunks) == 2
    assert len({chunk["chunk_id"] for chunk in chunks}) == len(chunks)
